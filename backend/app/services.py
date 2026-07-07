import io, math, re, logging, hashlib, shutil, threading, time
from datetime import datetime, timezone
from pathlib import Path
from .config import settings
from pypdf import PdfReader
from docx import Document as WordDocument
from openpyxl import load_workbook, Workbook
from sqlalchemy import select, delete, func
from sqlalchemy.orm import Session
from .models import Customer,Document,DocumentChunk,AuditLog,Questionnaire,Question,Answer,AnswerVersion,ProviderConfig,GlobalProviderConfig
from .providers import get_embeddings,get_llm,MANUAL
logger=logging.getLogger("questionnaire.rag")
logger.setLevel(logging.INFO)
def now():return datetime.now(timezone.utc).replace(tzinfo=None)
BACKUPS_TO_KEEP=10
def backup_sqlite_database(database_url=None,keep=BACKUPS_TO_KEEP):
    """Rolling snapshot of the local database; it holds all approved answers, so every backend start backs it up."""
    url=database_url or settings.database_url
    if not url.startswith("sqlite"):return None
    db_path=Path(url.split("///")[-1])
    if not db_path.exists():return None
    folder=db_path.parent/"backups";folder.mkdir(exist_ok=True)
    target=folder/f"{db_path.stem}-{datetime.now().strftime('%Y%m%d-%H%M%S-%f')}{db_path.suffix}"
    shutil.copy2(db_path,target)
    for old in sorted(folder.glob(f"{db_path.stem}-*{db_path.suffix}"))[:-keep]:old.unlink()
    logger.info("database_backup path=%s",target);return target
def clean_customer_answer(text,context):
    clean=text.strip()
    # Customer-ready answers are plain prose; strip markdown syntax that models emit despite instructions.
    clean=re.sub(r"\*\*(.+?)\*\*",r"\1",clean,flags=re.S);clean=re.sub(r"(?m)^#{1,6}\s+","",clean);clean=clean.replace("`","")
    clean=re.sub(r"(?m)^\s*[-*•]\s+","",clean)
    for item in context:clean=re.sub(re.escape(item["document"]),"",clean,flags=re.I)
    clean=re.sub(r"(?i)\b(?:source|document|chunk)\s*(?:id)?\s*[:#]?\s*\d+\b","",clean)
    clean=re.sub(r"(?i)\b(?:similarity|match)\s*(?:score)?\s*[:=]?\s*\d+(?:\.\d+)?%?","",clean)
    clean=re.sub(r"\[(?:source\s*)?\d+\]","",clean,flags=re.I)
    clean=re.sub(r"(?i)according to\s*(?:the)?\s*[,;:]?","",clean)
    clean=re.sub(r"\s+([,.;:])",r"\1",clean);clean=re.sub(r"[ \t]{2,}"," ",clean)
    return clean.strip(" ,;:-") or MANUAL

def parse_file(path:Path)->str:
    ext=path.suffix.lower()
    if ext==".pdf":return "\n".join(p.extract_text() or "" for p in PdfReader(path).pages)
    if ext==".docx":return "\n".join(p.text for p in WordDocument(path).paragraphs)
    if ext==".xlsx":
        wb=load_workbook(path,data_only=True,read_only=True)
        return "\n".join(" | ".join(str(v) for v in row if v is not None) for ws in wb for row in ws.iter_rows(values_only=True))
    return path.read_text(encoding="utf-8-sig",errors="replace")
def parse_pages(path:Path):
    if path.suffix.lower()==".pdf":return [(index+1,page.extract_text() or "") for index,page in enumerate(PdfReader(path).pages)]
    return [(None,parse_file(path))]
# NOTE: stripping repeated page headers before chunking was evaluated (scripts/evaluate_retrieval.py)
# and REJECTED: hit@1 dropped 0.90 -> 0.80 with no score gain — headers carry anchoring signal for the
# embedding model. Re-measure before reintroducing any chunk-cleaning idea.
def chunks(text,size=900,overlap=120):
    clean=re.sub(r"\s+"," ",text).strip();out=[];start=0
    while start<len(clean):
        end=min(len(clean),start+size);out.append(clean[start:end]);start=end-overlap if end<len(clean) else end
    return out
def config_for(db,cid):
    customer_cfg=db.scalar(select(ProviderConfig).where(ProviderConfig.customer_id==cid))
    if not customer_cfg:raise RuntimeError(f"Customer {cid} has no settings record")
    cfg=customer_cfg if customer_cfg.is_override else db.get(GlobalProviderConfig,1)
    if not cfg:raise RuntimeError("Global default settings are not configured")
    cfg.settings_source="customer_override" if customer_cfg.is_override else "global_default"
    if not (cfg.embedding_provider or "").strip():raise RuntimeError(f"Customer {cid} has no embedding provider selected")
    return cfg
def index_document(db,doc,text=None):
    cfg=config_for(db,doc.customer_id)
    try:
        doc.status="extracting";doc.error_message=None;db.commit();pages=[(None,text)] if text is not None else parse_pages(Path(doc.path));text="\n".join(value for _,value in pages);doc.extracted_text_length=len(text);logger.info("text_extracted customer=%s document=%s length=%s",doc.customer_id,doc.id,len(text))
        if not text.strip():raise ValueError("Text extraction produced no content")
        doc.status="chunking";db.commit();part_records=[(page,part) for page,value in pages for part in chunks(value,cfg.chunk_size,cfg.chunk_overlap)];parts=[part for _,part in part_records];logger.info("chunks_created customer=%s document=%s count=%s",doc.customer_id,doc.id,len(parts))
        if not parts:raise ValueError("No chunks were created")
        embedding_url=getattr(cfg,"embedding_base_url",None) or getattr(cfg,"base_url",None) or "<provider default>";logger.info("index_embedding_start customer=%s provider=%s model=%s base_url=%s chunks=%s settings_source=%s",doc.customer_id,cfg.embedding_provider,cfg.embedding_model,embedding_url,len(parts),cfg.settings_source)
        doc.status="embedding";db.commit();embeddings=get_embeddings(cfg).embed_batch(parts);logger.info("embeddings_created customer=%s provider=%s model=%s base_url=%s chunks=%s embeddings=%s settings_source=%s",doc.customer_id,cfg.embedding_provider,cfg.embedding_model,embedding_url,len(parts),len(embeddings),cfg.settings_source)
        if len(embeddings)!=len(parts) or any(not vector for vector in embeddings):raise ValueError(f"Expected {len(parts)} embeddings, received {len(embeddings)}")
        db.execute(delete(DocumentChunk).where(DocumentChunk.customer_id==doc.customer_id,DocumentChunk.document_id==doc.id))
        for i,((page_number,content),embedding) in enumerate(zip(part_records,embeddings)):db.add(DocumentChunk(customer_id=doc.customer_id,document_id=doc.id,content=content,embedding=embedding,ordinal=i,page_number=page_number))
        db.flush();stored=db.scalar(select(func.count(DocumentChunk.id)).where(DocumentChunk.customer_id==doc.customer_id,DocumentChunk.document_id==doc.id))
        if stored!=len(parts):raise ValueError(f"Expected {len(parts)} vector records, stored {stored}")
        dimension=len(embeddings[0]);doc.embedding_count=len(embeddings);doc.vector_count=stored;doc.embedding_provider=cfg.embedding_provider;doc.embedding_model=cfg.embedding_model;doc.embedding_dimension=dimension;doc.settings_source=cfg.settings_source;doc.indexed_chunk_size=cfg.chunk_size;doc.indexed_chunk_overlap=cfg.chunk_overlap;doc.last_indexed_at=now();doc.status="indexed";doc.error_message=None;db.commit();logger.info("vectors_stored customer=%s document=%s count=%s dimension=%s settings_source=%s",doc.customer_id,doc.id,stored,dimension,cfg.settings_source);return len(parts)
    except Exception as exc:
        db.rollback();failed=db.get(Document,doc.id);failed.status="failed";failed.error_message=str(exc);db.commit();logger.exception("index_failed customer=%s document=%s",doc.customer_id,doc.id);raise
def normalize_collections(collections):
    seen=[];
    for name in collections or []:
        clean=str(name).strip()
        if clean and clean not in seen:seen.append(clean)
    return seen
def delete_all_documents(db,cid):
    """Remove every document, chunk, and stored file for one customer; other tenants are untouched."""
    docs=list(db.scalars(select(Document).where(Document.customer_id==cid)))
    if not docs:return 0
    paths=[Path(d.path) for d in docs]
    db.execute(delete(DocumentChunk).where(DocumentChunk.customer_id==cid))
    db.execute(delete(Document).where(Document.customer_id==cid))
    db.add(AuditLog(customer_id=cid,action="documents_delete_all",entity_type="customer",entity_id=cid,details={"documents":len(paths)}));db.commit()
    for path in paths:path.unlink(missing_ok=True)
    logger.info("documents_delete_all customer=%s count=%s",cid,len(paths));return len(paths)
def ingest(db,path,name,cid,category="Company",collections=None):
    doc=Document(customer_id=cid,name=name,path=str(path),category=category,size_bytes=path.stat().st_size,status="uploaded",authority=authority_for(category),collections=normalize_collections(collections) or ["General"]);db.add(doc);db.commit();logger.info("document_uploaded customer=%s document=%s name=%s collections=%s",cid,doc.id,name,doc.collections)
    count=index_document(db,doc);db.add(AuditLog(customer_id=cid,action="document_upload",entity_type="document",entity_id=doc.id,details={"name":name,"chunks":count}));db.commit();return doc
def cosine(a,b):return sum(x*y for x,y in zip(a,b))/(math.sqrt(sum(x*x for x in a))*math.sqrt(sum(y*y for y in b)) or 1)
def relevance_terms(value):
    stop={"a","an","and","are","at","be","can","do","does","for","how","in","is","it","of","on","or","the","to","what","when","where","which","who","with","you","your","currently","provide","list"}
    return {token[:-1] if token.endswith("s") and len(token)>3 else token for token in re.findall(r"[a-z0-9]+",value.lower()) if token not in stop and len(token)>2}
def question_terms(value):
    aliases={"mfa":"multifactor","multi-factor":"multifactor","authentication":"auth","authenticate":"auth","available":"support","administrators":"admin","administrator":"admin","supported":"support","supporting":"support","encrypted":"encrypt","encryption":"encrypt","certified":"certification","certifications":"certification"}
    return {aliases.get(x,x) for x in relevance_terms(value.replace("multi-factor","multifactor"))}
def question_similarity(a,b):
    left,right=question_terms(a),question_terms(b)
    return len(left&right)/max(1,len(left|right))
def suggestion_pool(db,cid):
    """Prefetch every reuse candidate once — approved answers, their questions, and evidence-document state.
    Loading a 350-question review page previously re-scanned all approved answers per question (N+1)."""
    rows=db.execute(select(Answer,Question,Customer).join(Question,Question.id==Answer.question_id).join(Customer,Customer.id==Answer.customer_id).where(Answer.status=="approved",((Answer.customer_id==cid)|(Answer.global_approved==True)))).all()
    document_ids=set()
    for answer,_,_ in rows:document_ids.update(answer.evidence_document_ids or [])
    documents={d.id:d for d in db.scalars(select(Document).where(Document.id.in_(document_ids)))} if document_ids else {}
    pool=[]
    for answer,previous,customer in rows:
        evidence=[documents[i] for i in (answer.evidence_document_ids or []) if i in documents]
        evidence_current=bool(evidence) and all(d.enabled and d.status=="indexed" and d.customer_id==answer.customer_id and (not answer.approved_at or not d.updated_at or d.updated_at<=answer.approved_at) for d in evidence)
        pool.append({"answer":answer,"question_text":previous.text,"terms":question_terms(previous.text),"customer_name":customer.name,"evidence_current":evidence_current})
    return pool
def suggestions_from_pool(pool,question,cid,collections=None,category=None,exclude_answer_id=None,limit=3):
    selected=set(normalize_collections(collections))
    if not selected:return []
    terms=question_terms(question);found=[]
    for item in pool:
        answer=item["answer"]
        if answer.id==exclude_answer_id:continue
        # Reuse stays inside the selected Knowledge Collections unless an admin explicitly approved the answer globally.
        same_scope=answer.customer_id==cid and selected.intersection(answer.collections or []) and (not category or answer.category in {category,"General","Company"})
        if not same_scope and not answer.global_approved:continue
        score=len(terms&item["terms"])/max(1,len(terms|item["terms"]))
        # ISO string, not datetime: suggestions are embedded in the answer's debug_data JSON column, which cannot serialize datetimes.
        if score>=.28:found.append({"answer_id":answer.id,"question":item["question_text"],"answer":answer.text,"customer":item["customer_name"],"collections":answer.collections or [],"category":answer.category,"approved_at":answer.approved_at.isoformat() if answer.approved_at else None,"reviewer":answer.reviewer or "Reviewer","evidence":answer.sources,"evidence_status":"Current" if item["evidence_current"] else "Needs Review – evidence changed or unavailable","evidence_current":item["evidence_current"],"match_badge":"Global Approved Answer" if answer.global_approved else "Shared Collection Match","golden":answer.golden,"global_approved":answer.global_approved,"similarity":round(score,2)})
    return sorted(found,key=lambda x:(x["golden"],x["similarity"]),reverse=True)[:limit]
def approved_suggestions(db,question,cid,collections=None,category=None,exclude_answer_id=None,limit=3):
    return suggestions_from_pool(suggestion_pool(db,cid),question,cid,collections,category,exclude_answer_id,limit)
# Authority tiers 1-9: lower is more authoritative. 1-2 are reserved for Golden/Approved answers; documents start at 3.
CATEGORY_AUTHORITY={"Products":3,"Security":4,"Previous Questionnaires":5,"Company":6,"Support":7,"Operations":7,"Compliance":8,"Legal":8,"Marketing":9}
DEFAULT_AUTHORITY=6
PEER_AUTHORITY_GAP=1  # sources within this tier distance can genuinely contradict each other; wider gaps mean the higher tier supersedes
RELEVANT_SCORE=.25    # chunks below this retrieval score cannot veto or dilute an answer
RELIABLE_SCORE=.35
def authority_for(category):return CATEGORY_AUTHORITY.get(category,DEFAULT_AUTHORITY)
NEGATION_PATTERN=re.compile(r"\b(?:no|not|never|unsupported|does not|do not|cannot|isn't|is not|are not|aren't)\b")
def negated_subjects(text):
    # Terms within the clause around each negation — the only subjects a negative sentence actually denies.
    subjects=set()
    for match in NEGATION_PATTERN.finditer(text.lower()):subjects|=question_terms(text[max(0,match.start()-60):match.end()+60])
    return subjects
EXCLUSIVE_SENTINEL="__mutually_exclusive__"
# Words so common in questionnaire prose that sharing them proves nothing about the factual subject of a claim.
GENERIC_TERMS={"support","user","customer","platform","product","service","enterprise","company","solution","system","application","provide","offer","include"}
def conflict_subjects(left_text,right_text):
    """Terms of the factual subject one side denies while the other asserts; empty set when the claims are compatible."""
    left_text,right_text=left_text.lower(),right_text.lower()
    exclusive=lambda a,b:bool(re.search(r"\bcloud[- ]only\b",a)) and bool(re.search(r"\b(?:on[- ]?prem(?:ises)?|self[- ]hosted)[- ]only\b",b))
    if exclusive(left_text,right_text) or exclusive(right_text,left_text):return {EXCLUSIVE_SENTINEL}
    left_negated,right_negated=negated_subjects(left_text),negated_subjects(right_text)
    # One side must deny a subject the other asserts without negation: >=2 shared claim terms, at least one specific (non-generic).
    denied_by_left=(left_negated-right_negated)&(question_terms(right_text)-right_negated)
    denied_by_right=(right_negated-left_negated)&(question_terms(left_text)-left_negated)
    subjects=set()
    if len(denied_by_left)>=2 and denied_by_left-GENERIC_TERMS:subjects|=denied_by_left
    if len(denied_by_right)>=2 and denied_by_right-GENERIC_TERMS:subjects|=denied_by_right
    return subjects
def claims_conflict(left_text,right_text):return bool(conflict_subjects(left_text,right_text))
def llm_confirms_conflict(llm,left_text,right_text):
    """Precision filter: the lexical screen recalls candidate conflicts, the model confirms them. Fails closed."""
    try:
        verdict=llm.chat([{"role":"system","content":"You check factual consistency. Reply with exactly YES or NO."},{"role":"user","content":f"Statement A: {left_text}\n\nStatement B: {right_text}\n\nDo these statements make mutually incompatible factual claims about the same subject? Reply YES or NO."}])
        return "yes" in str(verdict).strip().lower()[:5]
    except Exception:return True
def analyze_evidence(context,llm=None,use_llm=False):
    """Rank retrieved evidence around the most authoritative relevant source instead of treating all pairs as potential conflicts."""
    relevant=[x for x in context if x.get("score",1)>=RELEVANT_SCORE]
    counts={"supporting":0,"complementary":0,"conflicting":0,"superseded":0,"unrelated":0}
    if not relevant:
        return {"primary":None,"roles":{},"counts":counts,"conflicting_documents":[],"superseded_documents":[],"consistency":0,"relevant_count":0}
    # Relevance leads, authority arbitrates: the primary is the most authoritative source among those near the top retrieval score.
    top=max(x.get("score",1) for x in relevant)
    band=[x for x in relevant if x.get("score",1)>=top-.15]
    primary=min(band,key=lambda x:(x.get("authority",DEFAULT_AUTHORITY),-x.get("score",1)))
    roles={primary["chunk_id"] if "chunk_id" in primary else id(primary):"primary"};conflicts=[];superseded=[]
    primary_terms=question_terms(primary["content"].lower())
    for item in relevant:
        if item is primary:continue
        key=item["chunk_id"] if "chunk_id" in item else id(item)
        item_terms=question_terms(item["content"].lower());shared=primary_terms&item_terms;similarity=len(shared)/max(1,len(primary_terms|item_terms))
        if claims_conflict(primary["content"],item["content"]):
            if item.get("authority",DEFAULT_AUTHORITY)-primary.get("authority",DEFAULT_AUTHORITY)>PEER_AUTHORITY_GAP:kind="superseded";superseded.append(item["document"])
            # A peer-authority conflict crashes the consistency score, so it must be confirmed, not just lexically suspected.
            elif not use_llm or llm is None or llm_confirms_conflict(llm,primary["content"],item["content"]):kind="conflicting";conflicts.extend([primary["document"],item["document"]])
            else:kind="complementary"
        elif similarity>=.68 or (len(item["content"])>40 and (item["content"].lower() in primary["content"].lower() or primary["content"].lower() in item["content"].lower())):kind="supporting"
        elif shared:kind="complementary"
        else:kind="unrelated"
        counts[kind]+=1;roles[key]=kind
    if counts["conflicting"]:consistency=.3
    else:consistency=round(min(.98,.9+.04*counts["supporting"]),2)
    return {"primary":primary,"roles":roles,"counts":counts,"conflicting_documents":sorted(set(conflicts)),"superseded_documents":sorted(set(superseded)),"consistency":consistency,"relevant_count":len(relevant)}
VERIFY_PROMOTION_MIN=.55  # borderline band: below this the evidence is too weak to promote even with model verification
def llm_supports_answer(llm,question,answer_text,context):
    """Second verification pass for borderline answers: promote to Ready only when the model confirms full evidentiary support."""
    evidence="\n\n".join(x["content"] for x in context)
    try:
        verdict=llm.chat([{"role":"system","content":"You verify that an answer is grounded in evidence. Reply with exactly YES or NO."},{"role":"user","content":f"Question: {question}\n\nAnswer: {answer_text}\n\nEvidence:\n{evidence}\n\nIs the answer fully supported by the evidence, with no invented facts? Reply YES or NO."}])
        return str(verdict).strip().lower().startswith("yes")
    except Exception:
        logger.exception("llm_verification_failed question=%r",question[:80]);return False
def verify_answer(answer_text,question,relevant_context,primary,llm=None,use_llm=False):
    """Generate-then-verify: check the drafted answer against relevant evidence instead of vetoing on chunk pairs."""
    contradictions=[];question_subject=question_terms(question)
    primary_authority=primary.get("authority",DEFAULT_AUTHORITY) if primary else DEFAULT_AUTHORITY
    for chunk in relevant_context:
        if primary and chunk is primary:continue
        if chunk.get("authority",DEFAULT_AUTHORITY)-primary_authority>PEER_AUTHORITY_GAP:continue  # lower tiers cannot veto the primary's answer
        subjects=conflict_subjects(answer_text,chunk["content"])
        if not subjects:continue
        # The contradiction must concern what the question actually asks; disagreements about side topics never block an answer.
        if EXCLUSIVE_SENTINEL not in subjects and not subjects&(question_subject-GENERIC_TERMS):continue
        # Lexical screen is the recall filter; the LLM is the precision filter.
        if not use_llm or llm is None or llm_confirms_conflict(llm,answer_text,chunk["content"]):contradictions.append(chunk["document"])
    return sorted(set(contradictions))
def retrieve(db,q,cid,limit=None,collections=None):
    cfg=config_for(db,cid);v=get_embeddings(cfg).embed_text(q);limit=limit or cfg.top_k;selected=set(normalize_collections(collections)) if collections is not None else None;logger.info("retrieval_query customer=%s query=%r top_k=%s dimension=%s collections=%s",cid,q,limit,len(v),sorted(selected) if selected is not None else "all")
    query=select(DocumentChunk,Document).join(Document,Document.id==DocumentChunk.document_id).where(DocumentChunk.customer_id==cid,Document.customer_id==cid,Document.status=="indexed",Document.enabled==True,Document.embedding_provider==cfg.embedding_provider,Document.embedding_model==cfg.embedding_model,Document.embedding_dimension==len(v))
    rows=db.execute(query).all()
    ranked=[];terms=relevance_terms(q)
    for c,d in rows:
        if len(c.embedding)!=len(v):continue
        # Retrieval isolation: only documents sharing at least one selected Knowledge Collection are searchable.
        if selected is not None and not selected.intersection(d.collections or []):continue
        # Hash embeddings are lexical test doubles; require a real token match to avoid collision false positives.
        content_terms=relevance_terms(c.content);overlap=len(terms.intersection(content_terms))/max(1,len(terms));name_overlap=len(terms.intersection(relevance_terms(d.name+" "+d.category)))/max(1,len(terms));vector_score=cosine(v,c.embedding)
        if cfg.embedding_provider=="mock" and overlap==0:continue
        ranked.append({"chunk_id":c.id,"content":c.content,"document":d.name,"document_id":d.id,"category":d.category,"collections":d.collections or [],"authority":d.authority or authority_for(d.category),"page_number":c.page_number,"score":min(1,vector_score+.2*overlap+.15*name_overlap),"vector_score":vector_score,"lexical_overlap":overlap,"embedding_provider":d.embedding_provider,"embedding_model":d.embedding_model})
    results=[];seen=set()
    # Authority breaks score ties so the more authoritative document leads the evidence list.
    for candidate in sorted(ranked,key=lambda x:(-x["score"],x["authority"])):
        if candidate["score"]<.05 or candidate["document_id"] in seen:continue
        results.append(candidate);seen.add(candidate["document_id"])
        if len(results)>=limit:break
    logger.info("retrieval_results customer=%s count=%s scores=%s",cid,len(results),[round(x["score"],4) for x in results]);return results
def detect_numbered_questions(text):
    """Distinct explicit Q-number markers in the source; used to warn when extraction loses questions."""
    return len(set(re.findall(r"(?im)^\s*q\s?(\d{1,4})[.):]",text)))
def build_questionnaire(db,path,name,cid,collections=None):
    cfg=config_for(db,cid);text=parse_file(path);llm=get_llm(cfg);qs=llm.extract_questions(text) or [x.strip() for x in text.splitlines() if x.strip()][:1000]
    detected=detect_numbered_questions(text)
    item=Questionnaire(customer_id=cid,name=name,path=str(path),collections=normalize_collections(collections));db.add(item);db.flush()
    for i,qtext in enumerate(qs):
        db.add(Question(customer_id=cid,questionnaire_id=item.id,text=qtext,ordinal=i))
    if detected>len(qs):logger.warning("question_extraction_mismatch customer=%s questionnaire=%s detected=%s extracted=%s",cid,item.id,detected,len(qs))
    db.add(AuditLog(customer_id=cid,action="questionnaire_upload",entity_type="questionnaire",entity_id=item.id,details={"name":name,"questions":len(qs),"numbered_detected":detected}));db.commit()
    item.detected_question_count=detected
    return item
ROLE_ORDER={"primary":0,"supporting":1,"complementary":2,"conflicting":3,"superseded":4,"unrelated":5,"additional":6}
def generate_one(db,q,cfg,llm,on_stage=None):
    stage=(lambda name:on_stage(name)) if on_stage else (lambda name:None)
    questionnaire=db.get(Questionnaire,q.questionnaire_id);collections=normalize_collections(questionnaire.collections);cache_key=hashlib.sha256(f"rag-v5|{cfg.embedding_provider}|{cfg.embedding_model}|{cfg.top_k}|{sorted(collections)}|{q.text}".encode()).hexdigest();cache_hit=q.retrieval_cache_key==cache_key and bool(q.retrieval_cache);stage("retrieving");ctx=q.retrieval_cache if cache_hit else (retrieve(db,q.text,q.customer_id,collections=collections) if collections else [])
    if q.retrieval_cache_key!=cache_key:q.retrieval_cache=ctx;q.retrieval_cache_key=cache_key
    analysis=analyze_evidence(ctx,llm,use_llm=cfg.llm_provider!="mock");primary=analysis["primary"];relevant=[x for x in ctx if x.get("score",1)>=RELEVANT_SCORE];roles=analysis["roles"];top_score=ctx[0]["score"] if ctx else 0
    existing=db.scalar(select(Answer).where(Answer.question_id==q.id,Answer.customer_id==q.customer_id));suggestions=approved_suggestions(db,q.text,q.customer_id,collections,exclude_answer_id=existing.id if existing else None)
    # Golden/approved reuse is validated against current documentation; a golden match with changed evidence is surfaced for review, never silently reused.
    validated=lambda x:x["evidence_current"] and (not (x["golden"] or x["global_approved"]) or bool(relevant and top_score>=RELIABLE_SCORE))
    preferred=next((x for x in suggestions if validated(x) and x["golden"] and x["similarity"]>=.55),None) or next((x for x in suggestions if validated(x) and x["similarity"]>=.82),None)
    stale_golden=None if preferred else next((x for x in suggestions if x["golden"] and x["similarity"]>=.55 and not x["evidence_current"]),None)
    started=time.perf_counter();generation_ctx=([primary]+[x for x in relevant if roles.get(x.get("chunk_id"),"")in{"supporting","complementary"}]) if primary else [];prompt=f"Question: {q.text}\n\nRetrieved context:\n"+"\n\n".join(x["content"] for x in generation_ctx);contradictions=[]
    if preferred:answer_text=preferred["answer"];logger.info("approved_answer_reused customer=%s question=%s source_answer=%s",q.customer_id,q.id,preferred["answer_id"])
    elif stale_golden:answer_text=stale_golden["answer"];logger.info("stale_golden_surfaced customer=%s question=%s source_answer=%s",q.customer_id,q.id,stale_golden["answer_id"])
    elif generation_ctx:
        # Generate first from the best evidence, then verify the draft against relevant peer evidence. A verified contradiction downgrades status but never erases the draft.
        stage("generating");logger.info("llm_called customer=%s question=%s context_count=%s",q.customer_id,q.id,len(generation_ctx));answer_text=clean_customer_answer(llm.generate_answer(q.text,generation_ctx,cfg.prompt_instructions),generation_ctx)
        if answer_text!=MANUAL:stage("verifying");contradictions=verify_answer(answer_text,q.text,relevant,primary,llm,use_llm=cfg.llm_provider!="mock")
    else:logger.info("llm_skipped customer=%s question=%s reason=no_relevant_evidence",q.customer_id,q.id);answer_text=MANUAL
    elapsed=round((time.perf_counter()-started)*1000,1)
    supporting=analysis["counts"]["supporting"];extra=supporting+analysis["counts"]["complementary"];retrieval_quality=round(min(1,top_score+(.1*ctx[1]["score"] if len(ctx)>1 else 0)),2) if ctx else 0;evidence_consistency=(analysis["consistency"] if not contradictions else .3) if relevant else 0
    authority_factor=max(.7,1-.05*(primary.get("authority",DEFAULT_AUTHORITY)-3)) if primary else 0
    provenance=(f"Answered from {primary['document']}" if primary else "")+(f", supported by {extra} additional source{'s' if extra!=1 else ''}." if primary and extra else "." if primary else "")
    if preferred:confidence=round(min(1,.6*preferred["similarity"]+.4),2)
    elif stale_golden:confidence=round(min(1,.5*stale_golden["similarity"]+.3),2)
    elif relevant:confidence=round(min(1,.55*retrieval_quality+.3*evidence_consistency+.15*authority_factor+.03*supporting),2)
    else:confidence=0
    # Borderline answers with reliable evidence get a second model pass; only a confirmed YES promotes them to Ready.
    llm_verified=False
    if (not preferred and not stale_golden and not contradictions and relevant and answer_text!=MANUAL
            and VERIFY_PROMOTION_MIN<=confidence<.7 and top_score>=RELIABLE_SCORE and cfg.llm_provider!="mock"):
        stage("verifying");llm_verified=llm_supports_answer(llm,q.text,answer_text,generation_ctx)
        if llm_verified:logger.info("llm_verification_promoted customer=%s question=%s confidence=%s",q.customer_id,q.id,confidence)
    stage("saving")
    answer=existing
    if preferred:status="approved_candidate";reason=("Golden Answer reused — validated against current documentation." if preferred["golden"] else "Previously approved answer reused — evidence verified as current.")
    elif stale_golden:status="needs_review";reason="Golden Answer matched, but source documentation changed since approval — verify before reuse."
    elif not relevant:status="manual_review";reason="No relevant documentation found in this product scope."
    elif answer_text==MANUAL:status="manual_review";reason="Current documentation does not support a grounded answer."
    elif contradictions:status="manual_review";reason=f"Documentation disagrees: {primary['document']} vs {', '.join(contradictions)} — reviewer decision required. Draft answer retained from {primary['document']}."
    elif (confidence>=.7 or llm_verified) and top_score>=RELIABLE_SCORE:status="approved_candidate";reason=provenance+(" Verified against source documentation." if llm_verified else "")
    else:status="needs_review";reason=provenance+(" Confirm wording before approval." if provenance else "Reviewer verification recommended.")
    current_sources=sorted([{"document":x["document"],"document_id":x["document_id"],"category":x["category"],"collections":x.get("collections",[]),"authority":x.get("authority",DEFAULT_AUTHORITY),"role":roles.get(x.get("chunk_id"),"additional"),"page_number":x.get("page_number"),"chunk_id":x["chunk_id"],"score":round(x["score"],4),"text_preview":x["content"][:700]} for x in ctx],key=lambda x:(ROLE_ORDER.get(x["role"],9),-x["score"]));sources=current_sources if preferred and preferred["global_approved"] else (preferred["evidence"] if preferred else current_sources)
    values={"text":answer_text,"confidence":confidence,"status":status,"sources":sources,"classification_reason":reason,"reused_from_answer_id":preferred["answer_id"] if preferred else (stale_golden["answer_id"] if stale_golden else None),"category":sources[0]["category"] if sources else "General","collections":collections,"evidence_document_ids":sorted(set(x.get("document_id") for x in sources if x.get("document_id"))),"debug_data":{"prompt":prompt,"retrieved_chunks":sources,"llm_response":answer_text,"execution_time_ms":elapsed,"cache_hit":cache_hit,"llm_verified":llm_verified,"conflicting_documents":contradictions,"superseded_documents":analysis["superseded_documents"],"evidence_analysis":{"counts":analysis["counts"],"consistency":analysis["consistency"],"relevant_count":analysis["relevant_count"],"primary_document":primary["document"] if primary else None},"retrieval_quality":retrieval_quality,"evidence_consistency":evidence_consistency,"answer_confidence":confidence,"authority_factor":authority_factor,"suggestions":suggestions}}
    if answer:
        for key,value in values.items():setattr(answer,key,value)
    else:answer=Answer(customer_id=q.customer_id,question_id=q.id,**values);db.add(answer)
    db.flush();db.add(AuditLog(customer_id=q.customer_id,action="previous_answer_reused" if preferred else "answer_generated",entity_type="answer",entity_id=answer.id,details={"source_answer_id":preferred["answer_id"] if preferred else None,"status":status,"reason":reason}));next_version=(db.scalar(select(func.max(AnswerVersion.version)).where(AnswerVersion.answer_id==answer.id)) or 0)+1;db.add(AnswerVersion(customer_id=q.customer_id,answer_id=answer.id,version=next_version,text=answer.text,confidence=answer.confidence,status=answer.status,sources=answer.sources));return answer
# Live generation progress, keyed by questionnaire id. In-memory by design: progress is ephemeral, answers are durable (committed per question).
GENERATION_PROGRESS={}
GENERATION_LOCK=threading.Lock()
GENERATION_STAGES={"preparing":"Preparing questionnaire","retrieving":"Retrieving evidence","generating":"Generating answer","verifying":"Verifying answer","saving":"Saving result","completed":"Completed"}
def new_progress(qid,cid):
    return {"questionnaire_id":qid,"customer_id":cid,"state":"running","stage":"preparing","total":0,"completed":0,"failed_count":0,"current_ordinal":0,"current_question":"","current_question_id":None,"question_status":{},"question_errors":{},"summary":None,"error":None,"cancel":False,"started_ts":time.time(),"finished_ts":None,"started_at":now().isoformat(),"finished_at":None}
def generation_progress(qid):return GENERATION_PROGRESS.get(qid)
def request_generation_cancel(qid):
    progress=GENERATION_PROGRESS.get(qid)
    if not progress or progress["state"]!="running":return False
    progress["cancel"]=True;return True
def start_generation(qid,cid,only_missing=False,include_approved=False,question_ids=None):
    with GENERATION_LOCK:
        current=GENERATION_PROGRESS.get(qid)
        if current and current["state"]=="running":return None
        progress=new_progress(qid,cid);GENERATION_PROGRESS[qid]=progress
    threading.Thread(target=_generation_worker,args=(qid,cid,progress,only_missing,include_approved,question_ids),daemon=True).start();return progress
def _generation_worker(qid,cid,progress,only_missing=False,include_approved=False,question_ids=None):
    from .db import SessionLocal
    try:
        with SessionLocal() as db:
            item=db.get(Questionnaire,qid)
            if not item or item.customer_id!=cid:raise RuntimeError("Questionnaire not found")
            run_generation(db,item,progress,only_missing=only_missing,include_approved=include_approved,question_ids=question_ids)
    except Exception as exc:
        progress.update(state="failed",error=str(exc),finished_ts=time.time(),finished_at=now().isoformat());logger.exception("generation_worker_failed questionnaire=%s",qid)
def run_generation(db,item,progress,on_question_complete=None,only_missing=False,include_approved=False,question_ids=None):
    """Generate every answer with per-question commits so completed work survives cancellation or a crash."""
    cfg=config_for(db,item.customer_id);llm=get_llm(cfg)
    questions=list(db.scalars(select(Question).where(Question.customer_id==item.customer_id,Question.questionnaire_id==item.id).order_by(Question.ordinal)))
    if question_ids is not None:
        # Explicit scope (filtered/retry/selected questions) overrides the derived filters; the caller chose deliberately.
        wanted=set(question_ids);questions=[q for q in questions if q.id in wanted]
    elif only_missing and questions:
        answered=set(db.scalars(select(Answer.question_id).where(Answer.customer_id==item.customer_id,Answer.question_id.in_([q.id for q in questions]))))
        questions=[q for q in questions if q.id not in answered]
    elif questions and not include_approved:
        # Approved answers are reviewed work; a full regeneration never replaces them unless explicitly requested.
        approved=set(db.scalars(select(Answer.question_id).where(Answer.customer_id==item.customer_id,Answer.status=="approved",Answer.question_id.in_([q.id for q in questions]))))
        questions=[q for q in questions if q.id not in approved]
    progress.update(total=len(questions),question_status={q.id:"queued" for q in questions},stage="preparing")
    item.status="generating";db.commit()
    for index,q in enumerate(questions,start=1):
        if progress.get("cancel"):progress["state"]="cancelled";break
        progress.update(current_ordinal=q.ordinal+1,current_question_id=q.id,current_question=q.text[:300]);progress["question_status"][q.id]="processing"
        try:
            answer=generate_one(db,q,cfg,llm,on_stage=lambda name:progress.__setitem__("stage",name))
            progress["stage"]="saving";db.commit()
            progress["question_status"][q.id]="manual_review" if answer.status=="manual_review" else "generated"
        except Exception as exc:
            db.rollback();progress["question_status"][q.id]="failed";progress["question_errors"][q.id]=str(exc)[:300] or exc.__class__.__name__;progress["failed_count"]+=1;logger.exception("question_generation_failed questionnaire=%s question=%s",item.id,q.id)
        progress["completed"]=index
        if on_question_complete:on_question_complete(progress)
    if progress["state"]=="cancelled":
        for question_id,status in list(progress["question_status"].items()):
            if status=="queued":progress["question_status"][question_id]="cancelled"
    item.status="generated"
    ready=db.scalar(select(func.count(Answer.id)).join(Question,Question.id==Answer.question_id).where(Question.questionnaire_id==item.id,Answer.status=="approved_candidate")) or 0
    check=db.scalar(select(func.count(Answer.id)).join(Question,Question.id==Answer.question_id).where(Question.questionnaire_id==item.id,Answer.status=="needs_review")) or 0
    manual=db.scalar(select(func.count(Answer.id)).join(Question,Question.id==Answer.question_id).where(Question.questionnaire_id==item.id,Answer.status=="manual_review")) or 0
    elapsed=round(time.time()-progress["started_ts"],1)
    progress["summary"]={"total":len(questions),"processed":progress["completed"],"ready":ready,"check":check,"manual":manual,"failed":progress["failed_count"],"elapsed_seconds":elapsed,"average_seconds":round(elapsed/progress["completed"],2) if progress["completed"] else 0}
    db.add(AuditLog(customer_id=item.customer_id,action="questionnaire_generation",entity_type="questionnaire",entity_id=item.id,details={"questions":len(questions),"processed":progress["completed"],"failed":progress["failed_count"],"cancelled":progress["state"]=="cancelled","elapsed_seconds":elapsed}));db.commit()
    if progress["state"]=="running":progress["state"]="completed"
    progress.update(stage="completed",current_question="",current_question_id=None,finished_ts=time.time(),finished_at=now().isoformat())
    return progress["completed"]
def generate_questionnaire(db,item):
    run_generation(db,item,new_progress(item.id,item.customer_id))
    return db.scalar(select(func.count(Question.id)).where(Question.questionnaire_id==item.id))
def _estimated_row_height(values,widths):
    """Excel does not auto-fit wrapped rows written by openpyxl; estimate the height from the longest cell."""
    lines=1
    for value,width in zip(values,widths):
        text=str(value) if value is not None else ""
        per_line=max(10,int(width*1.1))
        needed=sum(max(1,-(-len(part)//per_line)) for part in text.splitlines()) if text else 1
        lines=max(lines,needed)
    return min(300,lines*14+4)
def apply_version_restore(db,answer,version):
    """Restore a prior draft's text: the restored status comes from that version's own snapshot,
    not a fixed value — the reviewer sees the review state the answer actually had back then,
    still sitting in whichever triage bucket (Ready/Check/Manual/Approved) it belonged to."""
    snapshot=db.scalar(select(AnswerVersion).where(AnswerVersion.customer_id==answer.customer_id,AnswerVersion.answer_id==answer.id,AnswerVersion.version==version))
    if not snapshot:return None
    answer.text=snapshot.text;answer.confidence=snapshot.confidence;answer.status=snapshot.status;answer.sources=snapshot.sources
    next_version=(db.scalar(select(func.max(AnswerVersion.version)).where(AnswerVersion.answer_id==answer.id)) or 0)+1
    db.add(AnswerVersion(customer_id=answer.customer_id,answer_id=answer.id,version=next_version,text=answer.text,confidence=answer.confidence,status=answer.status,sources=answer.sources))
    return next_version
def export_xlsx(item,internal=False,customer_name=""):
    from openpyxl.styles import Alignment,Font,PatternFill
    from openpyxl.utils import get_column_letter
    wb=Workbook();ws=wb.active;ws.title="Internal Review" if internal else "Completed Questionnaire"
    headers=["Question","Final Approved Answer"]+(["Customer Scope","Knowledge Collections","Category","Evidence","Confidence","Reviewer","Approval Date","Answer Version","Golden Answer Used","Reuse Source Answer ID"] if internal else [])
    widths=[70,90]+([14,24,14,40,12,14,18,12,16,20] if internal else [])
    span=len(headers)
    ws.merge_cells(start_row=1,start_column=1,end_row=1,end_column=span)
    title=ws.cell(row=1,column=1,value=customer_name or "Customer Questionnaire");title.font=Font(bold=True,size=16);title.alignment=Alignment(horizontal="center",vertical="center");ws.row_dimensions[1].height=30
    ws.merge_cells(start_row=2,start_column=1,end_row=2,end_column=span)
    subtitle=ws.cell(row=2,column=1,value=f"{item.name} · {len(item.questions)} questions · exported {now().strftime('%d %b %Y')}");subtitle.font=Font(size=10,color="FF666666");subtitle.alignment=Alignment(horizontal="center")
    header_row=4
    for index,name in enumerate(headers,start=1):
        cell=ws.cell(row=header_row,column=index,value=name);cell.font=Font(bold=True,color="FFFFFFFF");cell.fill=PatternFill("solid",fgColor="FF176B63");cell.alignment=Alignment(vertical="center")
    ws.row_dimensions[header_row].height=20
    for index,width in enumerate(widths,start=1):ws.column_dimensions[get_column_letter(index)].width=width
    wrap=Alignment(wrap_text=True,vertical="top")
    row=header_row+1
    for q in sorted(item.questions,key=lambda x:x.ordinal):
        answer=q.answer
        if internal:values=[q.text,answer.text if answer else "",item.customer_id,", ".join((answer.collections if answer else item.collections) or []),answer.category if answer else "",", ".join(s["document"] for s in answer.sources) if answer else "",answer.confidence if answer else 0,answer.reviewer if answer else "",answer.approved_at if answer else None,"Latest",bool(answer and answer.golden),answer.reused_from_answer_id if answer else None]
        else:values=[q.text,answer.text if answer and answer.status=="approved" else ""]
        for index,value in enumerate(values,start=1):ws.cell(row=row,column=index,value=value).alignment=wrap
        ws.row_dimensions[row].height=_estimated_row_height(values[:2],widths[:2])
        row+=1
    ws.freeze_panes=ws.cell(row=header_row+1,column=1)
    stream=io.BytesIO();wb.save(stream);stream.seek(0);return stream
NUMBERED_QUESTION=re.compile(r"^\s*q?\s?\d{1,4}[.)]",re.I)
def export_pdf(item,customer_name=""):
    """Customer copy in the questionnaire's own reading order: each question with its approved answer beneath it."""
    from fpdf import FPDF
    pdf=FPDF(format="A4");pdf.set_margins(18,18,18);pdf.set_auto_page_break(auto=True,margin=18);pdf.add_page()
    font="helvetica";arial=Path("C:/Windows/Fonts/arial.ttf");arial_bold=Path("C:/Windows/Fonts/arialbd.ttf")
    if arial.exists():
        pdf.add_font("body","",str(arial));pdf.add_font("body","B",str(arial_bold if arial_bold.exists() else arial));font="body"
    clean=(lambda text:text) if font=="body" else (lambda text:text.encode("latin-1","replace").decode("latin-1"))
    def paragraph(height,text,align="L"):pdf.multi_cell(0,height,clean(text),align=align,new_x="LMARGIN",new_y="NEXT")
    pdf.set_font(font,"B",16);paragraph(9,customer_name or "Customer Questionnaire",align="C")
    pdf.set_font(font,"",10);pdf.set_text_color(110);paragraph(6,f"{item.name} · {len(item.questions)} questions · {now().strftime('%d %b %Y')}",align="C");pdf.set_text_color(0);pdf.ln(6)
    for index,q in enumerate(sorted(item.questions,key=lambda x:x.ordinal),start=1):
        label=q.text if NUMBERED_QUESTION.match(q.text) else f"{index}. {q.text}"
        pdf.set_font(font,"B",11);paragraph(6,label)
        answer=q.answer.text if q.answer and q.answer.status=="approved" else ""
        pdf.set_font(font,"",11)
        if answer:paragraph(6,answer)
        else:pdf.set_text_color(150);paragraph(6,"(no approved answer)");pdf.set_text_color(0)
        pdf.ln(3)
    return io.BytesIO(bytes(pdf.output()))
